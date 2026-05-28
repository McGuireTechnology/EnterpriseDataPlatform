from __future__ import annotations

import re
import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE_SETS = [
    {
        "source_dir": ROOT / "contrib" / "sources" / "policy-templates",
        "target_prefix": "",
        "source_label": None,
    },
    {
        "source_dir": ROOT / "contrib" / "sources" / "cis" / "Policy Templates",
        "target_prefix": "cis-",
        "source_label": "CIS Controls v8.1 Policy Template",
    },
]
TARGET_DIR = ROOT / "docs" / "grc" / "policy-templates"


def slugify(name: str) -> str:
    value = name.lower()
    value = re.sub(r"\s+\((\d+)\)$", r"-\1", value)
    value = re.sub(r"[^a-z0-9]+", "-", value)
    return value.strip("-")


def normalize_text(text: str) -> str:
    text = text.replace("\r", "\n").replace("\n", " ")
    text = text.replace("\xa0", " ")
    text = re.sub(r"<([^<>]+)>", r"[\1]", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def yaml_list(values: list[str]) -> list[str]:
    if not values:
        return ["control_mappings: []"]
    return ["control_mappings:"] + [f"  - {json.dumps(value)}" for value in values]


def cis_control_mappings(source: Path) -> list[str]:
    stem = source.stem
    controls: set[str] = set()
    for match in re.finditer(r"(?:^|[_\s])Controls[_\s]+((?:\d+)(?:[_\s,]+(?:\d+))*)", stem, flags=re.I):
        for number in re.findall(r"\d+", match.group(1)):
            controls.add(f"CIS Control {int(number)}")
    for match in re.finditer(r"(?:^|[_\s])Control[_\s]+(\d+)", stem, flags=re.I):
        controls.add(f"CIS Control {int(match.group(1))}")

    comparable_stem = stem.replace("_", " ").lower()
    if "account credential management" in comparable_stem:
        controls.add("CIS Control 6")
    if "security awareness training" in comparable_stem:
        controls.add("CIS Control 14")

    return sorted(controls, key=lambda value: int(re.search(r"\d+", value).group(0)))


def iter_blocks(parent: DocxDocument | _Cell):
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def paragraph_markdown(paragraph: Paragraph, first_content: bool) -> list[str]:
    text = normalize_text(paragraph.text)
    if not text:
        return []

    style = paragraph.style.name if paragraph.style is not None else ""
    heading_match = re.match(r"Heading\s+([1-6])$", style)
    section_match = re.match(r"SectionLvl([1-6])$", style)

    if style == "Title" or first_content:
        return [f"# {text}"]

    if heading_match:
        level = int(heading_match.group(1)) + 1
        return [f"{'#' * min(level, 6)} {text}"]

    if section_match:
        level = int(section_match.group(1)) + 1
        return [f"{'#' * min(level, 6)} {text}"]

    if ("List Bullet" in style or style == "List Paragraph") and not re.match(r"^[-*+]\s+", text):
        return [f"- {text}"]

    if ("List Number" in style or "List 1 Number" in style) and not re.match(r"^\d+[\.)]\s+", text):
        return [f"1. {text}"]

    return [text]


def escape_cell(text: str) -> str:
    text = normalize_text(text)
    text = text.replace("|", r"\|")
    return text or " "


def table_markdown(table: Table) -> list[str]:
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            parts = []
            for paragraph in cell.paragraphs:
                text = normalize_text(paragraph.text)
                if text:
                    parts.append(text)
            cells.append(escape_cell(" / ".join(parts)))
        if cells:
            rows.append(cells)

    if not rows:
        return []

    width = max(len(row) for row in rows)
    rows = [row + [" "] * (width - len(row)) for row in rows]
    header = rows[0]
    separator = ["---"] * width
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return lines


def title_from_doc(doc: DocxDocument, fallback: str) -> str:
    for paragraph in doc.paragraphs:
        text = normalize_text(paragraph.text)
        style = paragraph.style.name if paragraph.style is not None else ""
        if text and style in {"Title", "CoverTitle"}:
            return text
        if text and "Policy Template" in text and len(text) <= 100:
            return text
    return fallback


def convert_docx(source: Path, target_prefix: str = "", source_label: str | None = None) -> Path:
    doc = Document(source)
    fallback_title = re.sub(r"\s+\(\d+\)$", "", source.stem).replace("-", " ")
    fallback_title = fallback_title.replace("_", " ")
    fallback_title = re.sub(r"\s+", " ", fallback_title).strip()
    title = title_from_doc(doc, fallback_title)
    target = TARGET_DIR / f"{target_prefix}{slugify(source.stem)}.md"
    mappings = cis_control_mappings(source) if source_label else []

    lines = [
        "---",
        f"title: {json.dumps(title)}",
        f"template_source: {source.relative_to(ROOT).as_posix()}",
        "status: draft-template",
        "owner: TBD",
        "review_cadence: annual",
    ]
    if source_label:
        lines.append(f"source_label: {json.dumps(source_label)}")
    lines.extend(yaml_list(mappings))
    lines.extend(["---", "", f"# {title}", ""])

    has_content = False
    for block in iter_blocks(doc):
        block_lines: list[str]
        if isinstance(block, Paragraph):
            block_lines = paragraph_markdown(block, first_content=not has_content)
            if block_lines == [f"# {title}"]:
                has_content = True
                continue
        else:
            block_lines = table_markdown(block)

        if not block_lines:
            continue

        lines.extend(block_lines)
        lines.append("")
        has_content = True

    markdown = "\n".join(lines).rstrip() + "\n"
    TARGET_DIR.mkdir(parents=True, exist_ok=True)
    target.write_text(markdown, encoding="utf-8")
    return target


def main() -> None:
    total = 0
    for source_set in SOURCE_SETS:
        sources = sorted(source_set["source_dir"].glob("*.docx"))
        for source in sources:
            target = convert_docx(
                source,
                target_prefix=source_set["target_prefix"],
                source_label=source_set["source_label"],
            )
            print(f"{source.relative_to(ROOT)} -> {target.relative_to(ROOT)}")
        total += len(sources)
    print(f"Converted {total} policy templates.")


if __name__ == "__main__":
    main()
